"""Handlers for office document formats using :mod:`python-docx`."""

from __future__ import annotations

from pathlib import Path
from typing import Any, Dict

from .base_handler import BaseFileHandler

try:  # pragma: no-cover - optional dependency
    from docx import Document
except Exception:  # pragma: no-cover
    Document = None  # type: ignore


class OfficeHandler(BaseFileHandler):
    """Handle ``.docx`` files.

    The ``.doc`` and ``.odt`` extensions are recognised but not currently
    implemented and will raise a ``RuntimeError`` when used.
    """

    _extensions = {".doc", ".docx", ".odt"}

    def can_handle(self, file_path: str) -> bool:  # noqa: D401
        """See base class."""
        return Path(file_path).suffix.lower() in self._extensions

    def _require_docx(self) -> None:
        if Document is None:  # pragma: no cover
            raise RuntimeError("python-docx is required to handle Office files")

    def read_file(self, file_path: str) -> Dict[str, Any]:  # noqa: D401
        """See base class."""
        path = Path(file_path)
        ext = path.suffix.lower()
        if ext != ".docx":
            raise RuntimeError(f"Unsupported office format: {ext}")
        self._require_docx()
        try:
            doc = Document(str(path))
        except Exception as exc:  # pragma: no cover - corrupted file
            raise RuntimeError(f"Failed to read document: {exc}") from exc
        paragraphs = [p.text for p in doc.paragraphs]
        metadata = {
            "author": doc.core_properties.author,
            "created": doc.core_properties.created.isoformat()
            if doc.core_properties.created
            else None,
        }
        return {
            "title": path.stem,
            "content": "\n".join(paragraphs),
            "metadata": metadata,
            "structure": {"paragraphs": len(paragraphs)},
            "encoding": "utf-8",
            "format": ext.lstrip("."),
        }

    def save_file(self, file_path: str, data: Dict[str, Any]) -> bool:  # noqa: D401
        """See base class."""
        path = Path(file_path)
        ext = path.suffix.lower()
        if ext != ".docx":
            raise RuntimeError(f"Unsupported office format: {ext}")
        self._require_docx()
        try:
            doc = Document()
            for line in data.get("content", "").splitlines():
                doc.add_paragraph(line)
            doc.save(str(path))
        except Exception as exc:  # pragma: no cover - disk issues etc.
            raise RuntimeError(f"Failed to write document: {exc}") from exc
        return True
